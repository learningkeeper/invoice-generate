from docx import Document
import os
import datetime

def generate_doc(no, invoice_range):
    template_file_path = 'template-invoice.docx'
    output_file_path = f'Joe Yang Invoice-{no}.docx'
    variables = {
        "${NO}": no,
        "${INVOICE_RANGE}": invoice_range,
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    # setup your onboarding date here and range 
    onboarding = datetime.datetime.strptime("03/16/2022","%m/%d/%Y")
    for i in range(1,21):
        d = datetime.timedelta(days=14)
        a = onboarding.strftime("%m/%d/%Y")
        a1 = a[:-4] + a[-2:]
        onboarding = onboarding + d
        u1 = onboarding.strftime("%m/%d/%Y")
        u1 = u1[:-4] + u1[-2:]
        invoice_range = a1 + ' - ' +  u1
        no = str(i)
        generate_doc(no, invoice_range)


